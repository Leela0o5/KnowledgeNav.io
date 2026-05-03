from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import html2text
import pdfplumber  # type: ignore[import-untyped]
from docx import Document as DocxDocument  # type: ignore[import-untyped]


class UnsupportedFileTypeError(Exception):
    def __init__(self, suffix: str) -> None:
        super().__init__(f"Unsupported file type: {suffix}")
        self.suffix = suffix


@dataclass(frozen=True)
class Document:
    text: str
    source_file: str
    page_range: str | None


def _load_pdf(path: Path) -> list[Document]:
    docs: list[Document] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(text=text, source_file=str(path), page_range=str(i + 1)))
    return docs


def _load_markdown(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8")
    return [Document(text=text, source_file=str(path), page_range=None)]


def _load_html(path: Path) -> list[Document]:
    converter = html2text.HTML2Text()
    converter.ignore_links = False
    text = converter.handle(path.read_text(encoding="utf-8"))
    return [Document(text=text, source_file=str(path), page_range=None)]


def _load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(text=text, source_file=str(path), page_range=None)]


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8")
    return [Document(text=text, source_file=str(path), page_range=None)]


LOADER_REGISTRY: dict[str, Callable[[Path], list[Document]]] = {
    ".pdf": _load_pdf,
    ".md": _load_markdown,
    ".html": _load_html,
    ".htm": _load_html,
    ".docx": _load_docx,
    ".txt": _load_text,
}


def load_document(path: Path) -> list[Document]:
    loader = LOADER_REGISTRY.get(path.suffix.lower())
    if loader is None:
        raise UnsupportedFileTypeError(path.suffix)
    return loader(path)
